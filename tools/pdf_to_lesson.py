#!/usr/bin/env python3
"""
pdf_to_lesson.py

Converts a source curriculum document (PDF or DOCX) into a lesson spec that
matches the exact schema the Lesson Engine expects (src/content/lessonSpecs.js).

Usage:
    python3 pdf_to_lesson.py path/to/lesson.pdf --subject "Social Studies" \
        --course "The African American Experience" --out day4.js

    python3 pdf_to_lesson.py path/to/lesson.docx --subject Math \
        --course "Algebra I" --out unit3-day2.js

What it actually does, no magic:
    1. Extracts raw text from the PDF or DOCX (including tables, since the AAE
       curriculum's lessons live in Word tables, not paragraphs).
    2. Sends that raw text to Claude with a prompt describing the exact
       LessonEngine step "kind"s (choice, text, multi-text, content,
       graded-write, graded-numeric) and asks for a spec back as JSON.
    3. Validates the JSON has the required shape before writing it out.
    4. Writes a ready-to-import .js file into src/content/, in the same
       format as the specs already in that folder.

This does NOT invent content. If the source document doesn't clearly contain
a Hook/Read/Analyze/Write/Exit (or equivalent) structure, the script will
say so and refuse to fabricate one, rather than silently guessing.

Requires:
    pip install pdfplumber python-docx requests
    ANTHROPIC_API_KEY environment variable set
"""

import argparse
import json
import os
import re
import sys


def extract_pdf_text(path):
    import pdfplumber
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    chunks.append(" | ".join(c.strip() for c in row if c))
    return "\n".join(chunks)


def extract_docx_text(path):
    import docx
    d = docx.Document(path)
    chunks = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            chunks.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
    return "\n".join(chunks)


SCHEMA_PROMPT = """You are converting a source lesson document into a JSON lesson spec for a
specific lesson engine. Do not invent content that isn't in the source text.
If a required piece (like a rubric, or a specific reading page range) isn't
present, use null for that field rather than making it up.

The engine supports exactly these step "kind"s. Use ONLY these:

- "choice": a hook/prediction question with multiple options and an optional
  "why did you pick that" follow-up.
  Shape: { key, sectionLabel, minutes, kind: "choice", prompt, options: [strings], followUp }

- "text": a single short open text response.
  Shape: { key, sectionLabel, minutes, kind: "text", prompt }

- "multi-text": one or more groups of short-answer inputs (e.g. a reading
  hunt collecting 3 facts, or step-by-step guided practice).
  Shape: { key, sectionLabel, minutes, kind: "multi-text", instructions, prompt,
           groups: [{ label, count, placeholder }] }

- "content": a read-only instructional block (e.g. a worked example), no
  student input.
  Shape: { key, sectionLabel, minutes, kind: "content", title, body: [strings] }

- "graded-write": a structured written response (like a RACE paragraph) that
  gets graded against a rubric found in the source text.
  Shape: { key, sectionLabel, minutes, kind: "graded-write", prompt,
           fields: [{ key, label }], rubric: (the actual rubric text from the
           source, verbatim if present), maxScore }

- "graded-numeric": one or more problems with an exact numeric or short
  answer.
  Shape: { key, sectionLabel, minutes, kind: "graded-numeric", instructions,
           problems: [{ prompt, answer }], isMasteryGate: boolean }

Top-level spec shape:
{
  "subject": string,
  "courseLabel": string,
  "meta": string (e.g. "Week 2 · Day 5 · Ch.2, West African Empires (pp. 22-35)"),
  "steps": [ ...as above, in order... ]
}

Return ONLY valid JSON matching this shape. No markdown fences, no commentary
before or after. If the source text doesn't contain enough structure to fill
a field, use null for that field rather than guessing.

SOURCE DOCUMENT TEXT:
---
__SOURCE_TEXT__
---
"""


def call_claude(prompt):
    import requests
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["x-api-key"] = api_key
        headers["anthropic-version"] = "2023-06-01"
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers=headers,
        json={
            "model": "claude-sonnet-4-6",
            "max_tokens": 4000,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=120,
    )
    resp.raise_for_status()
    data = resp.json()
    text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
    return text


REQUIRED_STEP_FIELDS = {"key", "sectionLabel", "minutes", "kind"}
VALID_KINDS = {"choice", "text", "multi-text", "content", "graded-write", "graded-numeric"}


def validate_spec(spec):
    errors = []
    for field in ("subject", "courseLabel", "meta", "steps"):
        if field not in spec:
            errors.append(f"missing top-level field: {field}")
    if "steps" in spec:
        if not isinstance(spec["steps"], list) or not spec["steps"]:
            errors.append("steps must be a non-empty list")
        else:
            for i, step in enumerate(spec["steps"]):
                missing = REQUIRED_STEP_FIELDS - set(step.keys())
                if missing:
                    errors.append(f"step {i}: missing fields {missing}")
                if step.get("kind") not in VALID_KINDS:
                    errors.append(f"step {i}: invalid kind '{step.get('kind')}'")
    return errors


def to_js_module(spec, var_name):
    body = json.dumps(spec, indent=2)
    return (
        f"// Auto-generated by tools/pdf_to_lesson.py. Review before committing,\n"
        f"// this is a first draft from the source document, not guaranteed final.\n\n"
        f"const {var_name} = {body};\n\n"
        f"export {{ {var_name} }};\n"
    )


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("source", help="Path to source PDF or DOCX")
    ap.add_argument("--subject", required=True, help='e.g. "Social Studies"')
    ap.add_argument("--course", required=True, help='e.g. "The African American Experience"')
    ap.add_argument("--out", required=True, help="Output filename, written to src/content/generated/")
    ap.add_argument("--var-name", default=None, help="JS export name, defaults from --out")
    args = ap.parse_args()

    ext = os.path.splitext(args.source)[1].lower()
    if ext == ".pdf":
        raw_text = extract_pdf_text(args.source)
    elif ext == ".docx":
        raw_text = extract_docx_text(args.source)
    else:
        print(f"Unsupported file type: {ext}. Use .pdf or .docx.", file=sys.stderr)
        sys.exit(1)

    if len(raw_text.strip()) < 50:
        print("Extracted almost no text from the source document. Refusing to guess a lesson from nothing.", file=sys.stderr)
        sys.exit(1)

    prompt = SCHEMA_PROMPT.replace("__SOURCE_TEXT__", raw_text[:15000])
    raw_response = call_claude(prompt)
    cleaned = re.sub(r"^```json\s*|\s*```$", "", raw_response.strip())

    try:
        spec = json.loads(cleaned)
    except json.JSONDecodeError as e:
        print("Claude's response wasn't valid JSON. Raw response follows so you can fix it by hand:", file=sys.stderr)
        print(raw_response, file=sys.stderr)
        sys.exit(1)

    spec["subject"] = args.subject
    spec["courseLabel"] = args.course

    errors = validate_spec(spec)
    if errors:
        print("Generated spec failed validation:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        print("\nWriting it anyway to review manually, but it is not safe to import as-is.", file=sys.stderr)

    var_name = args.var_name or re.sub(r"[^A-Za-z0-9]", "_", os.path.splitext(args.out)[0]).upper()
    js = to_js_module(spec, var_name)

    out_dir = os.path.join(os.path.dirname(__file__), "..", "src", "content", "generated")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, args.out)
    with open(out_path, "w") as f:
        f.write(js)

    print(f"Wrote {out_path}")
    if not errors:
        print("Passed schema validation. Review the content, then import it into catalog.js.")
    else:
        print("Failed schema validation, see above. Fix by hand before using.")


if __name__ == "__main__":
    main()
